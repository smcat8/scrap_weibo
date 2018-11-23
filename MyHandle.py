#!/usr/bin/env python
#-*-coding:utf-8-*-

from HTMLParser import HTMLParser
from splinter.browser import Browser
import time
import splinter

import os
import re
import requests, urllib
import sys
import traceback
from datetime import datetime
from datetime import timedelta
from lxml import etree
import docx

#url = "https://passport.weibo.com/visitor/visitor?entry=miniblog&a=enter&url=https://weibo.com/ttarticle/p/show?id=2309404303470613106478&domain=.weibo.com&ua=php-sso_sdk_client-0.6.28&_rand=1541730661.0264"
#url = "http://t.cn/RYtILyJ"

#doc = docx.Document()



class MyHTMLParser(HTMLParser):
    def __init__(self, doc):
        HTMLParser.__init__(self)
        self.doc = doc
    #检索开头标签
    def handle_starttag(self,tag,attrs):
        print("Start tag:",tag)
        #匹配里面的项
        for attr in attrs:
            print("    attr:",attr)
        if 'img' == tag:
            for attr in attrs:
               if 'src' == attr[0]:
                   print(attr[1])
                   #将图片添加到Word文档中
                   urllib.request.urlretrieve(attr[1], '0.jpg')
                   self.doc.add_picture('0.jpg')
                   print("-----------Save img")

    #匹配结束标签
    def handle_endtag(self,tag):
        print("End tag  :",tag)
    #处理数据
    def handle_data(self,data):
        print("Data     :",data)
        #将每一段的内容添加到Word文档
        self.doc.add_paragraph(data)
        print("-----------Save into doc")
    #检索注释内容
    def handle_comment(self,data):
        print("Comment  :",data)
    #处理转义字符
#    def handle_entityref(self,name):
#        c = unichr(name2codepoint[name])
#        print("Named ent:",c)
    #处理转义的数字字符（ACSII）
    def handle_charref(self,name):
        if name.startswith('x'):
            c = unichr(int(name[1:],16))    #十六进制
        else:
            c = unichr(int(name))
        print("Num ent  :",c)
    #匹配HTML头
    def handle_decl(self,data):
        print("Decl     :",data)
        
    def add_paragraph(self,data):
        print("paragraph Data     :",data)
        self.doc.add_paragraph(data)
        print("-----------Save into doc")
        
    def add_heading(self,data):
        print("heading Data     :",data)
        self.doc.add_heading(data)
        print("-----------Save into doc")

class MyHandle:
    
    def __init__(self):
        self.browser = Browser("chrome")
    
    def __del__(self):
        self.browser.quit()
        
    def setParam(self, m_title, m_para, m_html):
        self.m_title = m_title
        self.m_para = m_para
        self.m_html = m_html
        
    def setContent(self, content):
        self.content = content

    def setParser(self, parser):
        self.parser = parser
    
    def handleURL(self, url):
        media_weibo_cn = [['h2'],['div[class="name m-text-cut"]','div[class="time"]', 'div[class="name m-text-cut"]'],['div[class="f-art"]']]
        keywords = {'div[class="m-feed"]':  media_weibo_cn}
        self.browser.visit(url)
        time.sleep(5)

        for key,value in keywords.items():
            content = self.browser.find_by_css(key)
            if content:
                self.setContent(content[0])
                self.setParam(value[0],value[1],value[2])
                self.handle()
            else:
                print("Can't find content!")
    
    def handle(self):
        if self.content and isinstance(self.content,splinter.driver.webdriver.WebDriverElement):
            if self.m_title and isinstance(self.m_title,list):
                for ti in self.m_title:
                    ti_content = self.content.find_by_css(ti)
                    #print(ti_content.text)
                    self.parser.add_heading(ti_content.text)
            if self.m_para and isinstance(self.m_para,list):
                for pa in self.m_para:
                    pa_content = self.content.find_by_css(pa)
                    #print(pa_content.text)
                    self.parser.add_paragraph(pa_content.text)
            if self.m_html and isinstance(self.m_html,list):
                for ht in self.m_html:
                    ht_content = self.content.find_by_css(ht)
                    self.parser.feed(ht_content.html)
        else:
            print("Wrong content!")



if __name__ == "__main__":
    #import pdb;pdb.set_trace()
    doc = docx.Document()
    parser = MyHTMLParser(doc)

    url = "https://media.weibo.cn/article?object_id=1022%3A2309404263284382559889&extparam=lmid--4263284376469708&luicode=10000011&lfid=1076032419394015&id=2309404263284382559889"

    myhandle = MyHandle()
    myhandle.setParser(parser)
    myhandle.handleURL(url)
    doc.save('Myhandle.docx')
