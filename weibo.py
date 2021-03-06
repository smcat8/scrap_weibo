#!/usr/bin/env python
# -*- coding: UTF-8 -*-

import os
import re
import requests, urllib
import sys
import traceback
from datetime import datetime
from datetime import timedelta
from lxml import etree
from MyHandle import MyHTMLParser, MyHandle
import docx
from docx.oxml.ns import qn
from docx.shared import Inches
import socket


class Weibo:
    cookie = {"Cookie": "YF-Page-G0=; SUB=; SUBP="}  # 将your cookie替换成自己的cookie

    # Weibo类初始化
    def __init__(self, user_id, filter=0):
        self.user_id = user_id  # 用户id，即需要我们输入的数字，如昵称为“Dear-迪丽热巴”的id为1669879400
        self.filter = filter  # 取值范围为0、1，程序默认值为0，代表要爬取用户的全部微博，1代表只爬取用户的原创微博
        self.username = ''  # 用户名，如“Dear-迪丽热巴”
        self.weibo_num = 0  # 用户全部微博数
        self.weibo_num2 = 0  # 爬取到的微博数
        self.following = 0  # 用户关注数
        self.followers = 0  # 用户粉丝数
        self.weibo_content = []  # 微博内容
        self.weibo_pic_link = []  # 原图
        self.weibo_link = []  # 转连接
        self.fake_link = [] # 无效连接
        self.weibo_place = []  # 微博位置
        self.publish_time = []  # 微博发布时间
        self.up_num = []  # 微博对应的点赞数
        self.retweet_num = []  # 微博对应的转发数
        self.comment_num = []  # 微博对应的评论数
        self.publish_tool = []  # 微博发布工具
        self.start_time = None
        self.end_time = None

    def setDoc(self, doc):
        self.doc = doc
    
    def setHandle(self, handle):
        self.handle = handle

    def setStartTime(self, start_time):
        self.start_time = start_time

    def setEndTime(self, end_time):
        self.end_time = end_time

    # 获取用户昵称
    def get_username(self):
        try:
            url = "https://weibo.cn/%d/info" % (self.user_id)
            html = requests.get(url, cookies=self.cookie).content
            selector = etree.HTML(html)
            username = selector.xpath("//title/text()")[0]
            self.username = username[:-3]
            print(u"用户名: " + self.username)
        except Exception as e:
            print("Error: ", e)
            traceback.print_exc()

    # 获取用户微博数、关注数、粉丝数
    def get_user_info(self):
        try:
            url = "https://weibo.cn/u/%d?filter=%d&page=1" % (
                self.user_id, self.filter)
            html = requests.get(url, cookies=self.cookie).content
            selector = etree.HTML(html)
            pattern = r"\d+\.?\d*"

            # 微博数
            str_wb = selector.xpath(
                "//div[@class='tip2']/span[@class='tc']/text()")[0]
            guid = re.findall(pattern, str_wb, re.S | re.M)
            for value in guid:
                num_wb = int(value)
                break
            self.weibo_num = num_wb
            print(u"微博数: " + str(self.weibo_num))

            # 关注数
            str_gz = selector.xpath("//div[@class='tip2']/a/text()")[0]
            guid = re.findall(pattern, str_gz, re.M)
            self.following = int(guid[0])
            print(u"关注数: " + str(self.following))

            # 粉丝数
            str_fs = selector.xpath("//div[@class='tip2']/a/text()")[1]
            guid = re.findall(pattern, str_fs, re.M)
            self.followers = int(guid[0])
            print(u"粉丝数: " + str(self.followers))
            print("===========================================================================")

        except Exception as e:
            print("Error: ", e)
            traceback.print_exc()

    # 获取"长微博"全部文字内容
    def get_long_weibo(self, weibo_link):
        try:
            html = requests.get(weibo_link, cookies=self.cookie).content
            selector = etree.HTML(html)
            info = selector.xpath("//div[@class='c']")[1]
            wb_content = info.xpath("div/span[@class='ctt']")[0].xpath(
                "string(.)").encode(sys.stdout.encoding, "ignore").decode(
                sys.stdout.encoding)
            wb_content = wb_content[1:]
            return wb_content
        except Exception as e:
            print("Error: ", e)
            traceback.print_exc()

    # 获取用户微博内容及对应的发布时间、点赞数、转发数、评论数
    def get_weibo_info(self):
        try:
            url = "https://weibo.cn/u/%d?filter=%d&page=1" % (
                self.user_id, self.filter)
            html = requests.get(url, cookies=self.cookie).content
            selector = etree.HTML(html)
            if selector.xpath("//input[@name='mp']") == []:
                page_num = 1
            else:
                page_num = (int)(selector.xpath(
                    "//input[@name='mp']")[0].attrib["value"])
            pattern = r"\d+\.?\d*"
            for page in range(1, page_num + 1):
                #if page > 5 : break
                url2 = "https://weibo.cn/u/%d?filter=%d&page=%d" % (
                    self.user_id, self.filter, page)
                html2 = requests.get(url2, cookies=self.cookie).content
                selector2 = etree.HTML(html2)
                info = selector2.xpath("//div[@class='c']")
                is_empty = info[0].xpath("div/span[@class='ctt']")
                if is_empty:
                    for i in range(0, len(info) - 2):
                        # 微博内容
                        str_t = info[i].xpath("div/span[@class='ctt']")
                        weibo_content = str_t[0].xpath("string(.)").encode(
                            sys.stdout.encoding, "ignore").decode(
                            sys.stdout.encoding)
                        weibo_content = weibo_content[:-1]
                        weibo_id = info[i].xpath("@id")[0][2:]
                        #html.unescape(urllib.parse.unquote('https://weibo.cn/sinaurl?f=w&amp;u=http%3A%2F%2Ft.cn%2FReBoAd9&amp;ep=GsPnbnhNK%2C2419394015%2CGsPnbnhNK%2C2419394015'))
                        #urllib.parse.parse_qs('https://weibo.cn/sinaurl?f=w&amp;u=http%3A%2F%2Ft.cn%2FReBoAd9&amp;ep=GsPnbnhNK%2C2419394015%2CGsPnbnhNK%2C2419394015')['u']
                        wb_link = []
                        a_link = info[i].xpath(
                            "div/span[@class='ctt']/a/@href")
                        if a_link:
                            if (a_link[-1] == "/comment/" + weibo_id or
                                    "/comment/" + weibo_id + "?" in a_link[-1]):
                                weibo_link = "https://weibo.cn" + a_link[-1]
                                wb_content = self.get_long_weibo(weibo_link)
                                if wb_content:
                                    weibo_content = wb_content
                            if 'u' in urllib.parse.parse_qs(a_link[-1]).keys():
                                wb_link.append(urllib.parse.parse_qs(a_link[-1])['u'][0])
                        if len(wb_link):
                            weibo_content += ' '.join(wb_link)
                        self.weibo_content.append(weibo_content)
                        self.weibo_link.append(' '.join(wb_link))
                        print(u"微博内容: " + weibo_content)

                        # 原图
                        wb_pic_link = []
                        p_link = info[i].xpath("div/a")
                        if p_link:
                            for p in p_link:
                                if p.xpath("text()") and u"原图" in p.xpath("text()")[0]:
                                    wb_pic_link.append(p.xpath("@href")[0])
                            print(u"原图: " + ' '.join(wb_pic_link))
                        self.weibo_pic_link.append(' '.join(wb_pic_link))

                        # 微博位置
                        div_first = info[i].xpath("div")[0]
                        a_list = div_first.xpath("a")
                        weibo_place = u"无"
                        for a in a_list:
                            if ("http://place.weibo.com/imgmap/center" in a.xpath("@href")[0] and
                                    a.xpath("text()")[0] == u"显示地图"):
                                weibo_place = div_first.xpath(
                                    "span[@class='ctt']/a")[-1]
                                if u"的秒拍视频" in div_first.xpath("span[@class='ctt']/a/text()")[-1]:
                                    weibo_place = div_first.xpath(
                                        "span[@class='ctt']/a")[-2]
                                weibo_place = weibo_place.xpath("string(.)").encode(
                                    sys.stdout.encoding, "ignore").decode(sys.stdout.encoding)
                                break
                        self.weibo_place.append(weibo_place)
                        print(u"微博位置: " + weibo_place)

                        # 微博发布时间
                        str_time = info[i].xpath("div/span[@class='ct']")
                        str_time = str_time[0].xpath("string(.)").encode(
                            sys.stdout.encoding, "ignore").decode(
                            sys.stdout.encoding)
                        publish_time = str_time.split(u'来自')[0]
                        if u"刚刚" in publish_time:
                            publish_time = datetime.now().strftime(
                                '%Y-%m-%d %H:%M')
                        elif u"分钟" in publish_time:
                            minute = publish_time[:publish_time.find(u"分钟")]
                            minute = timedelta(minutes=int(minute))
                            publish_time = (
                                datetime.now() - minute).strftime(
                                "%Y-%m-%d %H:%M")
                        elif u"今天" in publish_time:
                            today = datetime.now().strftime("%Y-%m-%d")
                            time = publish_time[3:]
                            publish_time = today + " " + time
                        elif u"月" in publish_time:
                            year = datetime.now().strftime("%Y")
                            month = publish_time[0:2]
                            day = publish_time[3:5]
                            time = publish_time[7:12]
                            publish_time = (
                                year + "-" + month + "-" + day + " " + time)
                        else:
                            publish_time = publish_time[:16]
                        self.publish_time.append(publish_time)
                        print(u"微博发布时间: " + publish_time)

                        # 微博发布工具
                        if len(str_time.split(u'来自')) > 1:
                            publish_tool = str_time.split(u'来自')[1]
                        else:
                            publish_tool = u"无"
                        self.publish_tool.append(publish_tool)
                        print(u"微博发布工具: " + publish_tool)

                        str_footer = info[i].xpath("div")[-1]
                        str_footer = str_footer.xpath("string(.)").encode(
                            sys.stdout.encoding, "ignore").decode(sys.stdout.encoding)
                        str_footer = str_footer[str_footer.rfind(u'赞'):]
                        guid = re.findall(pattern, str_footer, re.M)

                        # 点赞数
                        up_num = int(guid[0])
                        self.up_num.append(up_num)
                        print(u"点赞数: " + str(up_num))

                        # 转发数
                        retweet_num = int(guid[1])
                        self.retweet_num.append(retweet_num)
                        print(u"转发数: " + str(retweet_num))

                        # 评论数
                        comment_num = int(guid[2])
                        self.comment_num.append(comment_num)
                        print(u"评论数: " + str(comment_num))
                        print("===========================================================================")

                        self.weibo_num2 += 1

            if not self.filter:
                print(u"共" + str(self.weibo_num2) + u"条微博")
            else:
                print(u"共" + str(self.weibo_num) + u"条微博，其中" +
                       str(self.weibo_num2) + u"条为原创微博"
                       )
        except Exception as e:
            print("Error: ", e)
            traceback.print_exc()

    # 将爬取的信息写入文件
    def write_txt(self):
        try:
            if self.filter:
                result_header = u"\n\n原创微博内容: \n"
            else:
                result_header = u"\n\n微博内容: \n"
            result = (u"用户信息\n用户昵称：" + self.username +
                      u"\n用户id: " + str(self.user_id) +
                      u"\n微博数: " + str(self.weibo_num) +
                      u"\n关注数: " + str(self.following) +
                      u"\n粉丝数: " + str(self.followers) +
                      result_header
                      )
            for i in range(self.weibo_num2):
                text = (str(self.weibo_num2 - i) + ":" + self.weibo_content[i] + "\n" +
                        u"微博位置: " + self.weibo_place[i] + "\n" +
                        u"发布时间: " + self.publish_time[i] + "\n" +
                        u"点赞数: " + str(self.up_num[i]) +
                        u"   转发数: " + str(self.retweet_num[i]) +
                        u"   评论数: " + str(self.comment_num[i]) + "\n"
                        u"发布工具: " + self.publish_tool[i] + "\n\n"
                        )
                result = result + text
            file_dir = os.path.split(os.path.realpath(__file__))[
                0] + os.sep + "weibo"
            if not os.path.isdir(file_dir):
                os.mkdir(file_dir)
            file_path = file_dir + os.sep + "%d" % self.user_id + ".txt"
            f = open(file_path, "wb")
            f.write(result.encode(sys.stdout.encoding))
            f.close()
            print(u"微博写入文件完毕，保存路径:")
            print(file_path)
        except Exception as e:
            print("Error: ", e)
            traceback.print_exc()

    def auto_down(self, url, filename):
        try:
            urllib.request.urlretrieve(url,filename)
        except (socket.timeout, urllib.error.HTTPError) as e:
            count = 1
            while count <= 15:
                try:
                    urllib.request.urlretrieve(url, filename)
                    break
                except (socket.timeout, urllib.error.HTTPError) as e:
                    err_info = 'Reloading for %d time' % count if count == 1 else 'Reloading for %d times' % count
                    print(err_info)
                    count += 1
            if count > 15:
                print("下载失败")


    # 将爬取的信息写入文件
    def write_doc(self):
        try:
            requests.adapters.DEFAULT_RETRIES = 5
            if self.filter:
                result_header = u"\n\n原创微博内容: \n"
            else:
                result_header = u"\n\n微博内容: \n"
            result = (u"用户信息\n用户昵称：" + self.username +
                      u"\n用户id: " + str(self.user_id) +
                      u"\n微博数: " + str(self.weibo_num) +
                      result_header
                      )
            if self.doc:
                self.doc.add_heading(result)
            for i in range(self.weibo_num2-1,-1,-1):
                print(i)
                if self.start_time is not None:
                    if (self.start_time - datetime.strptime(self.publish_time[i].split()[0], "%Y-%m-%d")).days > 0:
                        continue
                if self.end_time is not None:
                    if (datetime.strptime(self.publish_time[i].split()[0], "%Y-%m-%d") - self.end_time).days > 0:
                        continue
                if self.doc:
                    self.doc.add_paragraph(str(self.weibo_num2 - i) + ":" + u"发布时间: " + self.publish_time[i])
                    self.doc.add_paragraph(self.weibo_content[i])
                    if "" != self.weibo_pic_link[i]:
                        for j in self.weibo_pic_link[i].split(' '):
                            print(j)
                            response = requests.get(j, cookies=self.cookie)
                            print(response.url)
                            pic = '1.jpg'
                            self.auto_down(response.url, pic)
                            try:
                                if (os.path.getsize(pic) > 40*1000):
                                    self.doc.add_picture(pic, width=Inches(5))
                                else:
                                    self.doc.add_picture(pic)
                            except:
                                print("Add picture fail:", response.url)
                    if "" != self.weibo_link[i]:
                        if self.publish_tool[i].find('weibo.com') > 0:
                            for j in self.weibo_link[i].split(' '):
                                response = requests.get(j, cookies=self.cookie)
                                print(self.publish_tool[i], response.url)
                                rs = self.handle.handleURL(response.url)
                                if not rs:
                                    self.fake_link.append(response.url)

            docName = self.username + ' ' + str(self.user_id)
            if self.start_time is not None:
                docName += "." + self.start_time.strftime("%Y-%m-%d")
            if self.end_time is not None:
                docName += "." + self.end_time.strftime("%Y-%m-%d")
            docName += ".docx" 
            self.doc.save(docName)      

        except Exception as e:
            print("Error: ", e)
            traceback.print_exc()
    # 运行爬虫
    def start(self):
        try:
            self.get_username()
            self.get_user_info()
            self.get_weibo_info()
            #self.write_txt()
            self.write_doc()
            print(u"信息抓取完毕")
            print(u"无效连接", self.fake_link)
            print("===========================================================================")
        except Exception as e:
            print("Error: ", e)


def main():
    if len(sys.argv) == 2:
        start_time = datetime.strptime(sys.argv[1], "%Y-%m-%d")
    elif len(sys.argv) == 3:
        start_time = datetime.strptime(sys.argv[1], "%Y-%m-%d")
        end_time = datetime.strptime(sys.argv[2], "%Y-%m-%d")
        if (start_time - end_time).days > 0:
            start_time, end_time = end_time, start_time
    try:
        doc = docx.Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        parser = MyHTMLParser(doc)
        myhandle = MyHandle()
        myhandle.setParser(parser)

        # 使用实例,输入一个用户id，所有信息都会存储在wb实例中
        user_id =  #  # 可以改成任意合法的用户id（爬虫的微博id除外）
        filter = 1  # 值为0表示爬取全部微博（原创微博+转发微博），值为1表示只爬取原创微博
        wb = Weibo(user_id, filter)  # 调用Weibo类，创建微博实例wb
        wb.setDoc(doc)
        wb.setHandle(myhandle)
        if 'start_time' in vars():
            wb.setStartTime(start_time)
        if 'end_time' in vars():
            wb.setEndTime(end_time)
        wb.start()  # 爬取微博信息
        print(u"用户名: " + wb.username)
        print(u"全部微博数: " + str(wb.weibo_num))
        print(u"关注数: " + str(wb.following))
        print(u"粉丝数: " + str(wb.followers))
        if wb.weibo_content:
            print(u"最新/置顶 微博为: " + wb.weibo_content[0])
            print(u"最新/置顶 微博位置: " + wb.weibo_place[0])
            print(u"最新/置顶 微博发布时间: " + wb.publish_time[0])
            print(u"最新/置顶 微博获得赞数: " + str(wb.up_num[0]))
            print(u"最新/置顶 微博获得转发数: " + str(wb.retweet_num[0]))
            print(u"最新/置顶 微博获得评论数: " + str(wb.comment_num[0]))
            print(u"最新/置顶 微博发布工具: " + wb.publish_tool[0])
    except Exception as e:
        print("Error: ", e)
        traceback.print_exc()


if __name__ == "__main__":
    #import pdb;pdb.set_trace()
    main()
